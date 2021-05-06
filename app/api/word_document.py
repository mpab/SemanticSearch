from document_utilities import DocUtil
from utilities import Folders, StringUtil
from context_types import SearchContext
from typing import Tuple
import docx
from docx import Document

def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

def word_document_create(search_context: SearchContext, filepath: str)-> Tuple[bool, str]:

    document = Document()
    # section = document.sections[0]
    # header = section.header
    # paragraph1 = header.paragraphs[0]
    # paragraph1.text = "\tIMAT 5314_1920_519: THESIS (DRAFT)"
    # paragraph2 = header.paragraphs[1]
    # paragraph2.text = "Abstract"

    heading = StringUtil.list_to_string(search_context.search_request.terms, ' ')
    document.add_heading('Search Context: (' + search_context.search_request.source_name + ') ' + heading, 0)
    
    document.add_heading('Index', level=1)
    para = document.add_paragraph()
    para.add_run('identifier_hash: ' + search_context.search_request.identifier_hash)
    para.add_run('\n') 

    ext = '.json'
    add_hyperlink(para, '../' + Folders.reference() + search_context.search_request.identifier_hash + ext, ext, 'blue', True)

    document.add_heading('References Table', level=1)

    # TODO: implement styles and/or use document template
    # https://pbpython.com/python-word-template.html

    #latent_styles = document.styles.latent_styles
    #latent_style = latent_styles.add_latent_style('Grid Table 5 Dark - Accent 1')
    # table = document.add_table(rows=1, cols=3, style = latent_style)

    table = document.add_table(rows=1, cols=5)
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Authors'
    hdr_cells[1].text = 'Publication'
    hdr_cells[2].text = 'Title'
    #hdr_cells[2].text = 'Abstract'
    #hdr_cells[2].text = 'Url'
    hdr_cells[3].text = 'Citations'
    hdr_cells[4].text = 'Link'
    #hdr_cells[4].text = 'Related'
    for page in search_context.search_result_pages:
        for result in page.search_results:
            row_cells = table.add_row().cells
            row_cells[0].text = result.authors
            row_cells[1].text = result.publication
            row_cells[2].text = result.title
            #row_cells[3].text = record.abstract
            #row_cells[2].text = record.url
            row_cells[3].text = str(result.citations)

            para = row_cells[4].add_paragraph()

            if result.pdf_status == 'validated':
                add_hyperlink(para, '../' + result.pdf_filepath, 'pdf', 'blue', True)
                para.add_run('\n') 
                add_hyperlink(para, result.site_url, 'url', 'green', True)
            elif result.site_url:
                add_hyperlink(para, result.site_url, 'url', 'green', True)
            else:
                row_cells[4].text = ''
   
    document.add_heading('References List', level=1)
    for page in search_context.search_result_pages:
        for result in page.search_results:
            ref_para = document.add_paragraph()
            lst = []
            lst.append(result.authors + '. ' + result.title)
            lst.append('Publication: ' + result.publication)
            if result.pdf_filetime:
                lst.append('Retrieved ' + result.pdf_filetime + ' from: ' + result.site_url)
            else:
                lst.append('Retrieved from: ' + result.site_url)

            lst.append('Abstract: "' + DocUtil.strip_newlines(result.abstract) + '"')
            ref_para.text = StringUtil.list_to_string(lst, '\n\t')

    document.save(filepath)

    return (True, 'created: ' + filepath)
    
